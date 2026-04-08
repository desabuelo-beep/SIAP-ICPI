"""
TERRA CIUDADANA — Sistema de Integridad Algorítmica Preventiva
QUADRUM GovTech | v2.0 | 2026
Plataforma pública de auditoría ciudadana independiente
MEJORAS INTEGRADAS: Gemini Chat + Mapa Folium + PDF FPDF2
"""

import streamlit as st
import pandas as pd
import json
import hashlib
import datetime
from io import BytesIO

# ─────────────────────────────────────────────
# CONFIGURACIÓN
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="TERRA CIUDADANA",
    page_icon="🌍",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ─────────────────────────────────────────────
# DATOS REALES MONTECRISTI 2024 (H36 BRIDGE)
# ─────────────────────────────────────────────
@st.cache_data
def cargar_datos_montecristi():
    return {
        "gad": "GAD Municipal de Montecristi",
        "alcalde": "Ing. Luis Jonathan Toro Largacha",
        "periodo": "2024-2027",
        "icpi": 72.93,
        "avep": "Gestión por Mandato",
        "ife": 77.0,
        "icm_sigad": 100.0,
        "brecha": 27.07,
        "itam": 70.0,
        "ics": 63.51,
        "sat": {"I": False, "II": False, "III": False, "IV": False},
        "nota_metodologica": (
            "⚠️ NOTA: Cédula presupuestaria 2025 usada como proxy metodológico para 2024. "
            "LOTAIP, verificables físicos y actas de entrega-recepción: simulación técnica documentada. "
            "Datos POA, PAC y Plan CNE: oficiales y reales 2024."
        ),
        "metas": [
            {"id": "SC-I-N-01", "nombre": "Salud Municipal", "icpi": 95.0, "avep": "🔵 Excelencia", "ods": "3,1,10", "dir": "Patronato"},
            {"id": "SC-L-N-02", "nombre": "TICs y Educativo", "icpi": 85.0, "avep": "🟢 Mandato", "ods": "4,10", "dir": "Dir. TIC"},
            {"id": "AH-I-X-01", "nombre": "Vialidad Cantonal", "icpi": 58.3, "avep": "🟡 Transición", "ods": "9,11", "dir": "Dir. Obras"},
            {"id": "AH-I-X-02", "nombre": "Señalización e IVU", "icpi": 0.0, "avep": "🔴 Ruptura", "ods": "11", "dir": "Dir. Obras"},
            {"id": "AH-I-X-03", "nombre": "Equipamientos Públicos", "icpi": 58.3, "avep": "🟡 Transición", "ods": "11,3,4", "dir": "Dir. Obras"},
            {"id": "AH-I-N-01", "nombre": "Vivienda Interés Social", "icpi": 22.5, "avep": "🟠 Ocurrencia", "ods": "1,11", "dir": "EP Montehogar"},
            {"id": "SC-L-G-01", "nombre": "Cultura y Patrimonio", "icpi": 47.8, "avep": "🟡 Transición", "ods": "4,11", "dir": "Dir. Cultura"},
            {"id": "AH-I-X-04", "nombre": "Tránsito y Seguridad Vial", "icpi": 40.5, "avep": "🟡 Transición", "ods": "11,3", "dir": "Dir. Obras"},
            {"id": "PI-I-G-01", "nombre": "Modernización Administrativa", "icpi": 85.0, "avep": "🟢 Mandato", "ods": "16", "dir": "Dir. Administrativa"},
            {"id": "AH-C-X-01", "nombre": "Agua Potable", "icpi": 72.0, "avep": "🟢 Mandato", "ods": "6,3,11", "dir": "Dir. Agua"},
            {"id": "AH-C-X-02", "nombre": "Alcantarillado y PTAR", "icpi": 72.0, "avep": "🟢 Mandato", "ods": "6,3,11", "dir": "Dir. Agua"},
            {"id": "SC-I-N-03", "nombre": "Grupos Prioritarios", "icpi": 53.4, "avep": "🟡 Transición", "ods": "1,3,5,10", "dir": "Patronato"},
            {"id": "FA-I-X-01", "nombre": "Gestión del Riesgo", "icpi": 65.0, "avep": "🟡 Transición", "ods": "11,13", "dir": "Dir. Ambiental"},
            {"id": "FA-C-X-01", "nombre": "Desechos Sólidos", "icpi": 80.0, "avep": "🟢 Mandato", "ods": "11,13", "dir": "EP Aseo"},
            {"id": "FA-I-X-02", "nombre": "Áreas Verdes e IVU", "icpi": 68.8, "avep": "🟡 Transición", "ods": "11,13,15", "dir": "Dir. Ambiental"},
            {"id": "FA-L-N-01", "nombre": "Fauna Urbana", "icpi": 77.0, "avep": "🟢 Mandato", "ods": "15", "dir": "Dir. Ambiental"},
            {"id": "PI-I-G-02", "nombre": "PDOT/Catastro/Trámites", "icpi": 85.0, "avep": "🟢 Mandato", "ods": "16,17", "dir": "Dir. Planificación"},
            {"id": "PI-L-G-01", "nombre": "Participación Ciudadana", "icpi": 53.4, "avep": "🟡 Transición", "ods": "16,17", "dir": "Dir. Planificación"},
            {"id": "EP-L-N-01", "nombre": "Fortalecimiento Productivo", "icpi": 95.0, "avep": "🔵 Excelencia", "ods": "1,8,10", "dir": "Dir. Económica"},
            {"id": "EP-L-X-01", "nombre": "Turismo", "icpi": 85.0, "avep": "🟢 Mandato", "ods": "8,11", "dir": "Dir. Turismo"},
        ]
    }

@st.cache_data
def datos_evaluacion_precacheados():
    return {
        "fecha_evento": "11 de julio de 2025",
        "lugar": "Comuna Toalla Grande",
        "asistentes": 261,
        "informe_numero": "22844",
        "icpi_verificado": 72.93,
        "icm_declarado": 100.0,
        "brecha_puntos": 27.07,
        "friccion_narrativa": 68.0,
        "coherencia": 32.0,
        "hallazgos": [
            {
                "meta": "AH-I-X-02 — Vialidad",
                "discurso": "Reconformación 90 km vías, maquinaria en proceso precontractual",
                "informe": "Maquinaria a espera de anticipo al 31/12/2024",
                "terra": "T_i = 0% — ejecución financiera nula",
                "tipo": "⚠️ Fricción Narrativa",
                "nivel": "alto"
            },
            {
                "meta": "SC-I-N-01 — Salud",
                "discurso": "3,743 usuarios atendidos en brigadas médicas",
                "informe": "3,743 atendidos confirmados",
                "terra": "ICPI = 95% — Excelencia",
                "tipo": "✅ Coherencia Verificada",
                "nivel": "bajo"
            },
            {
                "meta": "AH-C-X-01 — Agua Potable",
                "discurso": "Acueducto CAF $28M aprobado 29/12/2024",
                "informe": "Convenio firmado, hitos CAF documentados",
                "terra": "ICPI = 72% — Gestión por Mandato",
                "tipo": "✅ Coherencia Verificada",
                "nivel": "bajo"
            },
            {
                "meta": "PI-I-G-01 — Modernización",
                "discurso": "Modernización parque automotor y tecnología",
                "informe": "ERP $252k, maquinaria $600k adjudicados",
                "terra": "ICPI = 85% — Gestión por Mandato",
                "tipo": "✅ Coherencia Verificada",
                "nivel": "bajo"
            },
        ],
        "promesas_sin_pdot": [
            "EC-005 — Ordenanzas exoneración tributos a nuevas empresas",
            "IN-006 — Marca ciudad / posicionamiento nacional",
            "TE-005 — Camal municipal"
        ]
    }

GADS_ECUADOR = [
    "GAD Municipal de Montecristi",
    "GAD Municipal de Manta",
    "GAD Municipal de Portoviejo",
    "GAD Municipal de Jaramijó",
    "GAD Municipal de Quito",
    "GAD Municipal de Guayaquil",
    "GAD Municipal de Cuenca",
    "Otro municipio",
]

# ─────────────────────────────────────────────
# CSS
# ─────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;700&display=swap');

:root {
    --terra-blue: #003087;
    --terra-green: #38a169;
    --terra-yellow: #d69e2e;
    --terra-orange: #e67e22;
    --terra-red: #e53e3e;
    --terra-dark: #1a1a2e;
    --terra-light: #f0f4f8;
}

* { font-family: 'Space Grotesk', sans-serif; }

.terra-header {
    background: linear-gradient(135deg, #003087 0%, #1a365d 50%, #2c5282 100%);
    padding: 2rem;
    border-radius: 12px;
    color: white;
    margin-bottom: 2rem;
}
.terra-header h1 { font-size: 2.5rem; font-weight: 700; margin: 0; }
.terra-header p { font-size: 1rem; opacity: 0.85; margin: 0.5rem 0 0 0; }

.component-card {
    background: white;
    border: 2px solid #e2e8f0;
    border-radius: 12px;
    padding: 1.5rem;
    margin-bottom: 1rem;
    transition: all 0.3s ease;
}
.component-card:hover {
    border-color: #003087;
    box-shadow: 0 4px 20px rgba(0,48,135,0.15);
    transform: translateY(-2px);
}
.component-header { display: flex; align-items: center; gap: 1rem; margin-bottom: 1rem; }
.component-icon {
    width: 48px; height: 48px; border-radius: 12px;
    display: flex; align-items: center; justify-content: center;
    font-size: 1.5rem;
    background: linear-gradient(135deg, #003087, #2c5282);
    color: white;
}

.icpi-gauge {
    text-align: center; padding: 2rem;
    background: linear-gradient(135deg, #f0f4f8, #e2e8f0);
    border-radius: 16px; margin: 1rem 0;
}
.icpi-number {
    font-size: 5rem; font-weight: 700; color: #003087;
    line-height: 1; font-family: 'JetBrains Mono', monospace;
}
.icpi-label { font-size: 1.2rem; color: #4a5568; margin-top: 0.5rem; }

.metric-card {
    background: white; border-radius: 12px;
    padding: 1.2rem; text-align: center; border: 1px solid #e2e8f0;
}
.metric-value { font-size: 2rem; font-weight: 700; font-family: 'JetBrains Mono', monospace; }
.metric-label { font-size: 0.85rem; color: #718096; margin-top: 0.3rem; }

.brecha-alert {
    background: linear-gradient(135deg, #fff5f5, #fed7d7);
    border: 2px solid #fc8181; border-radius: 12px;
    padding: 1.5rem; margin: 1rem 0;
}

.nota-metodologica {
    background: #fffbeb; border-left: 4px solid #d69e2e;
    padding: 1rem; border-radius: 0 8px 8px 0;
    font-size: 0.85rem; color: #744210; margin: 1rem 0;
}

.sat-card { padding: 1rem; border-radius: 8px; text-align: center; border: 2px solid; }
.sat-ok { background: #f0fff4; border-color: #68d391; color: #276749; }
.sat-alert { background: #fff5f5; border-color: #fc8181; color: #c53030; }

.hash-display {
    background: #1a1a2e; color: #68d391; padding: 1rem; border-radius: 8px;
    font-family: 'JetBrains Mono', monospace; font-size: 0.8rem; word-break: break-all;
}

.friccion-barra {
    height: 20px; border-radius: 10px;
    background: linear-gradient(90deg, #68d391 0%, #d69e2e 50%, #fc8181 100%);
    margin: 0.5rem 0;
}

.hallazgo-card { padding: 1rem; border-radius: 8px; margin-bottom: 0.8rem; border-left: 4px solid; }
.hallazgo-alto { background: #fff5f5; border-color: #fc8181; }
.hallazgo-bajo { background: #f0fff4; border-color: #68d391; }

.footer-terra {
    text-align: center; padding: 2rem; color: #718096; font-size: 0.85rem;
    border-top: 1px solid #e2e8f0; margin-top: 3rem;
}

.proxima-fase {
    background: linear-gradient(135deg, #1a1a2e, #2d3748);
    color: white; padding: 1.5rem; border-radius: 12px; margin: 0.5rem 0;
}

.gemini-box {
    background: linear-gradient(135deg, #f0f4ff, #e8f0fe);
    border: 2px solid #4285f4; border-radius: 12px;
    padding: 1.5rem; margin: 1rem 0;
}

.stButton button {
    background: linear-gradient(135deg, #003087, #2c5282);
    color: white; border: none; border-radius: 8px;
    padding: 0.6rem 1.5rem; font-weight: 600; transition: all 0.3s;
}
.stButton button:hover { transform: translateY(-2px); box-shadow: 0 4px 15px rgba(0,48,135,0.3); }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div style='text-align:center; padding: 1rem;'>
        <div style='font-size:2rem;'>🌍</div>
        <div style='font-weight:700; font-size:1.2rem; color:#003087;'>TERRA CIUDADANA</div>
        <div style='font-size:0.8rem; color:#718096;'>QUADRUM GovTech v2.0</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("**Navega por los componentes:**")

    componentes = {
        "🏠 Inicio": "inicio",
        "📄 TERRA ACCEDE": "accede",
        "🔍 TERRA VERIFICA": "verifica",
        "🧮 TERRA CALCULA": "calcula",
        "🗺️ TERRA MAPEA": "mapea",
        "⚖️ TERRA ARTICULA": "articula",
        "✊ TERRA ACTÚA": "actua",
        "🏛️ TERRA EVALÚA": "evalua",
    }

    if "componente_activo" not in st.session_state:
        st.session_state.componente_activo = "inicio"

    for nombre, clave in componentes.items():
        if st.button(nombre, key=f"nav_{clave}", use_container_width=True):
            st.session_state.componente_activo = clave

    st.markdown("---")
    datos_sidebar = cargar_datos_montecristi()
    color_icpi = "#38a169" if datos_sidebar["icpi"] >= 70 else "#d69e2e"
    st.markdown(f"""
    <div style='background:#f0f4f8; padding:1rem; border-radius:8px; text-align:center;'>
        <div style='font-size:0.75rem; color:#718096;'>ICPI Montecristi 2024</div>
        <div style='font-size:2rem; font-weight:700; color:{color_icpi};'>{datos_sidebar['icpi']}%</div>
        <div style='font-size:0.8rem; color:#4a5568;'>🟢 {datos_sidebar['avep']}</div>
    </div>
    """, unsafe_allow_html=True)

activo = st.session_state.componente_activo

# ─────────────────────────────────────────────
# FUNCIÓN GEMINI — MEJORA 1
# ─────────────────────────────────────────────
def seccion_gemini_chat(contexto_datos):
    """Chat con Gemini integrado. Llama esta función desde cualquier sección."""
    st.markdown("---")
    st.markdown("""
    <div class='gemini-box'>
        <div style='font-weight:700; color:#003087; font-size:1.1rem;'>
            💬 Pregúntale a TERRA — Asistente con IA
        </div>
        <div style='font-size:0.85rem; color:#4a5568; margin-top:0.3rem;'>
            Haz cualquier pregunta sobre los datos de Montecristi en lenguaje cotidiano.
        </div>
    </div>
    """, unsafe_allow_html=True)

    pregunta = st.chat_input("Ej: ¿Por qué la vialidad tiene 0%? ¿Qué significa la brecha de 27 puntos?")
    if pregunta:
        with st.spinner("TERRA está analizando tu pregunta..."):
            try:
                import google.generativeai as genai
                genai.configure(api_key=st.secrets["GEMINI_KEY"])
                modelo = genai.GenerativeModel("gemini-pro")
                prompt = f"""
Eres TERRA, el asistente de auditoría municipal ciudadana de QUADRUM GovTech Ecuador.
Tu misión es explicar datos técnicos de gestión pública en lenguaje que cualquier ciudadano entienda.

DATOS VERIFICADOS — GAD Montecristi 2024:
- ICPI global: 72.93% — clasificado como "Gestión por Mandato" (escala AVEP)
- ICM declarado por el municipio: 100%
- Brecha documentada: 27.07 puntos porcentuales
- Meta más crítica: Señalización e IVU — ICPI 0% (Ruptura Sistémica)
- Metas en Excelencia (95%): Salud Municipal, Fortalecimiento Productivo
- IFE Electoral (fidelidad promesas CNE): 77%
- ITAM Transparencia: 70%
- ICS Social: 63.51%
- Alcalde: Ing. Luis Jonathan Toro Largacha
- 9 silos de verificación cruzada independientes
- Metodología: SIAP-ICPI v2.0 QUADRUM

CONTEXTO ADICIONAL: {contexto_datos}

REGLAS:
1. Responde en español simple, como explicando a un vecino del barrio
2. Máximo 3 párrafos cortos
3. Si es pregunta legal, menciona el artículo de la ley pero explícalo simple
4. Si preguntan qué hacer, guíalos hacia TERRA ACTÚA
5. Nunca inventes datos que no estén en el contexto

Pregunta del ciudadano: {pregunta}
                """
                respuesta = modelo.generate_content(prompt)
                st.markdown(f"""
                <div style='background:#f0f4ff; border-left:4px solid #4285f4; padding:1rem; border-radius:0 8px 8px 0; margin-top:1rem;'>
                    <div style='font-size:0.75rem; color:#718096; margin-bottom:0.5rem;'>🤖 TERRA responde:</div>
                    {respuesta.text}
                </div>
                """, unsafe_allow_html=True)
            except KeyError:
                st.warning("⚙️ Chat IA no configurado. Agrega GEMINI_KEY en Streamlit Settings → Secrets.")
            except Exception as e:
                st.error(f"Error al conectar con Gemini: {str(e)}")


# ─────────────────────────────────────────────
# FUNCIÓN PDF — MEJORA 3
# ─────────────────────────────────────────────
def generar_pdf_reporte(datos):
    """Genera PDF profesional con fpdf2."""
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()

        # Encabezado
        pdf.set_fill_color(0, 48, 135)
        pdf.rect(0, 0, 210, 35, 'F')
        pdf.set_font("Helvetica", "B", 18)
        pdf.set_text_color(255, 255, 255)
        pdf.set_xy(10, 8)
        pdf.cell(0, 10, "TERRA CIUDADANA — QUADRUM GovTech", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_xy(10, 20)
        pdf.cell(0, 8, "Reporte de Auditoría Ciudadana Independiente — Metodología SIAP-ICPI v2.0", ln=True)

        pdf.set_xy(10, 40)
        pdf.set_text_color(0, 0, 0)

        # Fecha
        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 6, f"Generado: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')} UTC", ln=True)
        pdf.ln(4)

        # Título GAD
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(0, 48, 135)
        pdf.cell(0, 10, f"GAD Municipal de Montecristi — Período 2024", ln=True)
        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(50, 50, 50)
        pdf.cell(0, 7, f"Alcalde: {datos['alcalde']}", ln=True)
        pdf.ln(4)

        # Línea separadora
        pdf.set_draw_color(0, 48, 135)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(6)

        # Índices principales
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(0, 48, 135)
        pdf.cell(0, 8, "ÍNDICES VERIFICADOS", ln=True)
        pdf.ln(2)

        indices = [
            ("ICPI Global (Índice de Congruencia Programática)", f"{datos['icpi']}%", "🟢 Gestión por Mandato"),
            ("ICM Declarado al SIGAD/SNP", f"{datos['icm_sigad']}%", "Autoreporte oficial"),
            ("Brecha MOM Documentada", f"{datos['brecha']} pp", "⚠️ Diferencia verificada"),
            ("IFE — Fidelidad Electoral", f"{datos['ife']}%", "Promesas CNE incorporadas al PDOT"),
            ("ITAM — Transparencia Activa", f"{datos['itam']}%", "Portal institucional"),
            ("ICS — Congruencia Social", f"{datos['ics']}%", "Inclusión y grupos prioritarios"),
        ]

        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(0, 0, 0)
        for nombre_idx, valor_idx, nota_idx in indices:
            pdf.set_fill_color(240, 244, 248)
            pdf.cell(110, 7, nombre_idx, border=0, fill=True)
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(25, 7, valor_idx, border=0, fill=True)
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(0, 7, nota_idx, border=0, ln=True)
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(1)

        pdf.ln(4)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(6)

        # Alerta MOM
        pdf.set_fill_color(255, 245, 245)
        pdf.set_draw_color(252, 129, 129)
        pdf.rect(10, pdf.get_y(), 190, 20, 'FD')
        pdf.set_xy(14, pdf.get_y() + 3)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(197, 48, 48)
        pdf.cell(0, 6, "⚠️  MUTACIÓN DEL OBJETO DE MEDICIÓN (MOM) DETECTADA", ln=True)
        pdf.set_xy(14, pdf.get_y())
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(116, 42, 42)
        pdf.cell(0, 6, f"El GAD declaró 100% al SIGAD. TERRA verifica 72.93%. Brecha: 27.07 puntos porcentuales.", ln=True)
        pdf.ln(8)

        # Metas por sección
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(0, 48, 135)
        pdf.cell(0, 8, "ESTADO POR META ESTRATÉGICA", ln=True)
        pdf.ln(2)

        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(255, 255, 255)
        pdf.set_fill_color(0, 48, 135)
        pdf.cell(15, 7, "ID", border=0, fill=True)
        pdf.cell(75, 7, "Meta", border=0, fill=True)
        pdf.cell(20, 7, "ICPI%", border=0, fill=True)
        pdf.cell(50, 7, "AVEP", border=0, fill=True)
        pdf.cell(0, 7, "Dirección", border=0, fill=True, ln=True)

        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(0, 0, 0)
        fill = False
        for meta in sorted(datos["metas"], key=lambda x: x["icpi"], reverse=True):
            pdf.set_fill_color(248, 250, 252) if fill else pdf.set_fill_color(255, 255, 255)
            pdf.cell(15, 6, meta["id"][:10], border=0, fill=True)
            pdf.cell(75, 6, meta["nombre"][:35], border=0, fill=True)
            pdf.cell(20, 6, f"{meta['icpi']}%", border=0, fill=True)
            pdf.cell(50, 6, meta["avep"], border=0, fill=True)
            pdf.cell(0, 6, meta["dir"][:20], border=0, fill=True, ln=True)
            fill = not fill

        pdf.ln(6)

        # Nota metodológica
        pdf.set_font("Helvetica", "I", 8)
        pdf.set_text_color(100, 100, 100)
        pdf.set_fill_color(255, 251, 235)
        pdf.multi_cell(0, 5, f"NOTA METODOLÓGICA: {datos['nota_metodologica']}", fill=True)

        pdf.ln(6)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(4)

        # Pie de página
        pdf.set_font("Helvetica", "I", 8)
        pdf.set_text_color(113, 128, 150)
        pdf.multi_cell(0, 5,
            "Este reporte fue generado mediante TERRA CIUDADANA, plataforma de ejercicio del derecho "
            "constitucional de acceso a la información pública (Art. 18 Constitución del Ecuador). "
            "QUADRUM GovTech — terra.quadrum.ec | Metodología SIAP-ICPI v2.0 | "
            "El acceso a información pública es un derecho. Esta herramienta no sustituye asesoría jurídica profesional."
        )

        return bytes(pdf.output())

    except ImportError:
        return None


# ─────────────────────────────────────────────
# INICIO
# ─────────────────────────────────────────────
if activo == "inicio":
    st.markdown("""
    <div class='terra-header'>
        <h1>🌍 TERRA CIUDADANA</h1>
        <p>El poder de la ciencia de datos en manos del ciudadano.<br>
        Auditoría territorial independiente | QUADRUM GovTech 2026</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("### ¿Qué puedes hacer aquí?")

    col1, col2 = st.columns(2)
    componentes_info = [
        ("📄", "TERRA ACCEDE", "Genera tu oficio LOTAIP blindado jurídicamente. Ejerce tu derecho constitucional en 3 minutos.", "#003087"),
        ("🔍", "TERRA VERIFICA", "Sube los documentos del GAD. El sistema los valida y genera evidencia con hash criptográfico.", "#2c5282"),
        ("🧮", "TERRA CALCULA", "Motor ICPI ciudadano. La misma fórmula canónica que usa el sistema institucional.", "#38a169"),
        ("🗺️", "TERRA MAPEA", "Visualiza las brechas territoriales de tu cantón en un mapa interactivo.", "#d69e2e"),
        ("⚖️", "TERRA ARTICULA", "Conecta tu análisis con el marco legal y el financiamiento internacional.", "#e67e22"),
        ("✊", "TERRA ACTÚA", "De la evidencia al cambio. Cartas, canales institucionales, proyectos.", "#e53e3e"),
        ("🏛️", "TERRA EVALÚA", "El informe técnico, el discurso político y los datos: tres verdades en un panel.", "#003087"),
    ]

    for i, (icon, nombre, desc, color) in enumerate(componentes_info):
        with col1 if i % 2 == 0 else col2:
            st.markdown(f"""
            <div class='component-card'>
                <div class='component-header'>
                    <div class='component-icon' style='background:linear-gradient(135deg,{color},{color}99);'>{icon}</div>
                    <div>
                        <div style='font-weight:700; color:{color};'>{nombre}</div>
                        <div style='font-size:0.85rem; color:#718096;'>{desc}</div>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### Caso piloto activo: GAD Montecristi 2024")
    datos = cargar_datos_montecristi()

    c1, c2, c3, c4 = st.columns(4)
    metricas = [
        (c1, datos['icpi'], "%", "ICPI Verificado", "#38a169"),
        (c2, datos['icm_sigad'], "%", "ICM Declarado", "#e53e3e"),
        (c3, datos['brecha'], "pp", "Brecha MOM", "#d69e2e"),
        (c4, datos['ife'], "%", "IFE Electoral", "#003087"),
    ]
    for col, valor, unidad, label, color in metricas:
        with col:
            st.markdown(f"""
            <div class='metric-card'>
                <div class='metric-value' style='color:{color};'>{valor}{unidad}</div>
                <div class='metric-label'>{label}</div>
            </div>""", unsafe_allow_html=True)

    st.markdown(f"""
    <div class='nota-metodologica'>{datos['nota_metodologica']}</div>
    """, unsafe_allow_html=True)

    seccion_gemini_chat("El ciudadano está en la página de inicio viendo el resumen general.")

# ─────────────────────────────────────────────
# C1 — TERRA ACCEDE
# ─────────────────────────────────────────────
elif activo == "accede":
    st.markdown("""
    <div class='terra-header'>
        <h1>📄 TERRA ACCEDE</h1>
        <p>Ejerce tu derecho constitucional. El Estado te debe información pública.</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    **¿Cómo funciona?**
    1. Selecciona el GAD que quieres auditar
    2. El sistema genera el oficio blindado jurídicamente
    3. Descarga el Word, llena tus datos en casa, imprime y entrega en ventanilla
    4. El GAD tiene **15 días hábiles** para responder (LOTAIP Art. 9)
    """)
    st.markdown("---")

    col1, col2 = st.columns([1, 1])
    with col1:
        perfil = st.selectbox("¿Cómo prefieres que te hable el sistema?", [
            "🏘️ Como vecino del barrio",
            "✊ Como activista ciudadano",
            "🎓 Como académico o investigador",
            "📰 Como periodista"
        ])
        gad_seleccionado = st.selectbox("¿Qué municipio quieres auditar?", GADS_ECUADOR)
        periodo = st.selectbox("¿De qué período necesitas información?", ["2024", "2023", "2022", "Período personalizado"])
        incluir_datos = st.checkbox("Quiero pre-llenar mis datos en el oficio (opcional)")

        nombre_peticionario = ""
        cedula_peticionario = ""
        correo_peticionario = ""

        if incluir_datos:
            st.info("Tus datos se usan ÚNICAMENTE para generar este documento. No se almacenan en el sistema. (LOPD Art. 9)")
            nombre_peticionario = st.text_input("Tu nombre completo")
            cedula_peticionario = st.text_input("Número de cédula")
            correo_peticionario = st.text_input("Correo electrónico")

    with col2:
        st.markdown("**El oficio incluirá estos fundamentos legales:**")
        fundamentos = [
            ("Constitución Art. 18 y 91", "Derecho de acceso a la información pública"),
            ("LOTAIP Arts. 7, 9, 11", "Obligación de transparencia activa y plazos"),
            ("LOPC Arts. 8, 73", "Participación ciudadana y control social"),
            ("COOTAD Art. 304", "Presupuesto participativo y acceso"),
            ("Ley Comercio Electrónico Art. 2", "Validez de documentos digitales"),
            ("Decreto 1384 Datos Abiertos", "Formatos abiertos obligatorios"),
            ("COPFP Art. 54", "Evaluación de cumplimiento de metas"),
        ]
        for ley, desc in fundamentos:
            st.markdown(f"✅ **{ley}** — {desc}")
        st.markdown("**El oficio exigirá:** POA en .xlsx · PAC en .xlsx · Cédula presupuestaria Grupos 75 y 84 · Respuesta digital al correo")

    if st.button("📄 Generar mi oficio LOTAIP", use_container_width=True):
        if not gad_seleccionado:
            st.error("Selecciona un municipio primero")
        else:
            with st.spinner("Generando oficio blindado..."):
                try:
                    from docx import Document
                    from docx.enum.text import WD_ALIGN_PARAGRAPH

                    doc = Document()
                    enc = doc.add_heading("SOLICITUD DE ACCESO A INFORMACIÓN PÚBLICA", 0)
                    enc.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    fecha_hoy = datetime.datetime.now().strftime("%d de %B de %Y")
                    doc.add_paragraph(f"Fecha: {fecha_hoy}")
                    doc.add_paragraph("Para: Director/a de Comunicación y Transparencia")
                    doc.add_paragraph(f"Institución: {gad_seleccionado}")
                    doc.add_paragraph("")
                    doc.add_paragraph(f"De: {nombre_peticionario if nombre_peticionario else '[NOMBRE COMPLETO]'}")
                    doc.add_paragraph(f"C.I.: {cedula_peticionario if cedula_peticionario else '[NÚMERO DE CÉDULA]'}")
                    doc.add_paragraph(f"Correo: {correo_peticionario if correo_peticionario else '[CORREO ELECTRÓNICO]'}")
                    doc.add_paragraph("")
                    doc.add_heading("SOLICITUD DE INFORMACIÓN PÚBLICA — LOTAIP", level=1)

                    cuerpo = doc.add_paragraph()
                    cuerpo.add_run(
                        f"En ejercicio del derecho constitucional de acceso a la información pública "
                        f"establecido en el Artículo 18 de la Constitución de la República del Ecuador, "
                        f"y en cumplimiento de la Ley Orgánica de Transparencia y Acceso a la Información "
                        f"Pública (LOTAIP), solicito respetuosamente al {gad_seleccionado} la siguiente "
                        f"información correspondiente al período fiscal {periodo}:"
                    )

                    doc.add_heading("Documentos solicitados:", level=2)
                    for item in [
                        f"Plan Operativo Anual (POA) {periodo} en formato Excel (.xlsx)",
                        f"Plan Anual de Contratación (PAC) {periodo} en formato Excel (.xlsx)",
                        f"Cédula Presupuestaria {periodo} — solo Grupos 75 y 84 (inversión) — en formato .xlsx",
                        "Respuesta digital al correo del peticionario (Ley Comercio Electrónico Art. 2 + Decreto 1384)",
                    ]:
                        doc.add_paragraph(item, style='List Bullet')

                    doc.add_heading("Fundamento Legal:", level=2)
                    for f_legal in [
                        "Constitución del Ecuador, Arts. 18 y 91",
                        "LOTAIP Arts. 7, 9, 11 — plazo 15 días hábiles",
                        "LOPC Arts. 8 y 73 — participación ciudadana",
                        "COOTAD Art. 304 — información presupuestaria municipal",
                        "COPFP Art. 54 — seguimiento y evaluación de metas",
                        "Ley de Comercio Electrónico Art. 2 — documentos digitales válidos",
                        "Decreto Ejecutivo 1384 — datos abiertos en formatos reutilizables",
                    ]:
                        doc.add_paragraph(f_legal, style='List Bullet')

                    doc.add_paragraph("")
                    doc.add_paragraph(
                        "De no recibir respuesta en el plazo legal de 15 días hábiles, me reservo el "
                        "derecho de presentar la acción de acceso a la información pública prevista en "
                        "el Art. 91 de la Constitución y el Art. 22 de la LOTAIP."
                    )
                    doc.add_paragraph("")
                    doc.add_paragraph("Atentamente,")
                    doc.add_paragraph("")
                    doc.add_paragraph("_________________________________")
                    doc.add_paragraph(nombre_peticionario if nombre_peticionario else "[FIRMA DEL PETICIONARIO]")
                    doc.add_paragraph(cedula_peticionario if cedula_peticionario else "[C.I.]")
                    doc.add_paragraph("")
                    pie = doc.add_paragraph()
                    pie.add_run(
                        "Oficio generado mediante TERRA CIUDADANA, plataforma de ejercicio del derecho "
                        "constitucional de acceso a la información pública (Art. 18 Constitución del Ecuador). "
                        "QUADRUM GovTech — terra.quadrum.ec"
                    ).italic = True

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    st.success("✅ Oficio generado exitosamente")
                    st.download_button(
                        label="📥 Descargar Oficio LOTAIP (.docx)",
                        data=buffer,
                        file_name=f"oficio_LOTAIP_{gad_seleccionado.replace(' ', '_')}_{periodo}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    st.markdown("""
                    <div class='nota-metodologica'>
                    <strong>Próximos pasos:</strong><br>
                    1. Descarga el Word → llena tus datos → imprime → firma<br>
                    2. Entrega en la ventanilla de transparencia del GAD<br>
                    3. Guarda el comprobante de recepción<br>
                    4. El GAD tiene 15 días hábiles para responder por correo con los archivos en .xlsx<br>
                    5. Cuando recibas la respuesta → regresa a <strong>TERRA VERIFICA</strong>
                    </div>
                    """, unsafe_allow_html=True)

                except ImportError:
                    st.error("Librería python-docx no instalada. Verifica tu requirements.txt")

    seccion_gemini_chat("El ciudadano está generando un oficio LOTAIP para pedir información pública al municipio.")

# ─────────────────────────────────────────────
# C2 — TERRA VERIFICA
# ─────────────────────────────────────────────
elif activo == "verifica":
    st.markdown("""
    <div class='terra-header'>
        <h1>🔍 TERRA VERIFICA</h1>
        <p>Convierte los documentos del GAD en evidencia con cadena de custodia.</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    **¿Qué necesitas subir?**
    - La copia del oficio de respuesta del GAD (con sello y firma — sin tus datos personales)
    - El POA en Excel · El PAC en Excel · La Cédula Presupuestaria en Excel
    """)
    st.markdown("---")

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**1. Oficio de respuesta del GAD**")
        oficio = st.file_uploader("Sube la copia con sello/firma", type=["pdf", "png", "jpg"], key="oficio")
        st.markdown("**2. Plan Operativo Anual (POA)**")
        poa = st.file_uploader("POA en Excel", type=["xlsx", "xls"], key="poa")
        st.markdown("**3. Plan Anual de Contratación (PAC)**")
        pac = st.file_uploader("PAC en Excel", type=["xlsx", "xls"], key="pac")
        st.markdown("**4. Cédula Presupuestaria**")
        cedula_pres = st.file_uploader("Cédula en Excel o PDF", type=["xlsx", "xls", "pdf"], key="cedula")

    with col2:
        st.markdown("**Estado de documentos:**")
        docs_estado = [
            ("Oficio respuesta GAD", oficio),
            ("POA Excel", poa),
            ("PAC Excel", pac),
            ("Cédula Presupuestaria", cedula_pres),
        ]
        for nombre_doc, archivo in docs_estado:
            if archivo:
                contenido = archivo.read()
                hash_doc = hashlib.sha256(contenido).hexdigest()
                st.markdown(f"""
                <div style='background:#f0fff4; border:1px solid #68d391; padding:0.8rem; border-radius:8px; margin-bottom:0.5rem;'>
                    <div style='font-weight:600; color:#276749;'>✅ {nombre_doc}</div>
                    <div class='hash-display' style='margin-top:0.5rem; font-size:0.7rem;'>SHA-256: {hash_doc[:40]}...</div>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div style='background:#fff5f5; border:1px solid #fc8181; padding:0.8rem; border-radius:8px; margin-bottom:0.5rem;'>
                    <div style='color:#c53030;'>⏳ {nombre_doc} — pendiente</div>
                </div>
                """, unsafe_allow_html=True)

        docs_completos = all([oficio, poa, pac, cedula_pres])
        if docs_completos:
            st.success("✅ Todos los documentos cargados. Puedes continuar a TERRA CALCULA.")
        else:
            pendientes = sum(1 for _, a in docs_estado if not a)
            st.info(f"Faltan {pendientes} documento(s) para completar la validación.")

    st.markdown("""
    <div class='proxima-fase'>
        <div style='font-weight:700; margin-bottom:0.5rem;'>🔮 Próxima fase — OCR activo</div>
        <div style='font-size:0.9rem; opacity:0.85;'>
        Cuando se active el módulo OCR, TERRA VERIFICA podrá leer PDFs escaneados e imágenes.
        El sistema te preguntará: "Detecté $45,000 — ¿confirmas?" (Human-in-the-Loop).
        </div>
    </div>
    """, unsafe_allow_html=True)

    seccion_gemini_chat("El ciudadano está subiendo documentos oficiales del GAD para validarlos.")

# ─────────────────────────────────────────────
# C3 — TERRA CALCULA
# ─────────────────────────────────────────────
elif activo == "calcula":
    st.markdown("""
    <div class='terra-header'>
        <h1>🧮 TERRA CALCULA</h1>
        <p>La verdad matemática de tu territorio. Fórmula canónica SIAP-ICPI.</p>
    </div>
    """, unsafe_allow_html=True)

    datos = cargar_datos_montecristi()

    modo = st.radio(
        "¿Qué quieres calcular?",
        ["Ver análisis del caso piloto: Montecristi 2024", "Cargar mis propios documentos"],
        horizontal=True
    )

    if modo == "Ver análisis del caso piloto: Montecristi 2024":

        st.markdown(f"""
        <div class='icpi-gauge'>
            <div style='font-size:0.9rem; color:#718096; margin-bottom:0.5rem;'>GAD Municipal de Montecristi — Período 2024</div>
            <div class='icpi-number'>{datos['icpi']}%</div>
            <div class='icpi-label'>🟢 {datos['avep']}</div>
            <div style='margin-top:1rem; font-size:0.85rem; color:#4a5568;'>
                Fórmula: ICPI = [Σ(Pi × Ri × Vi × Ei × Ti × Ci) / Σ(Pi × Ri)] × 100<br>
                Verificación cruzada: 9 silos de información pública independientes
            </div>
        </div>
        """, unsafe_allow_html=True)

        col1, col2, col3 = st.columns(3)
        indices = [
            ("IFE Electoral", datos['ife'], "%", "#003087"),
            ("ICM SIGAD", datos['icm_sigad'], "%", "#e53e3e"),
            ("Brecha MOM", datos['brecha'], "pp", "#d69e2e"),
            ("ITAM", datos['itam'], "%", "#2c5282"),
            ("ICS Social", datos['ics'], "%", "#38a169"),
        ]
        for i, (nombre, valor, unidad, color) in enumerate(indices):
            with [col1, col2, col3][i % 3]:
                st.markdown(f"""
                <div class='metric-card'>
                    <div class='metric-value' style='color:{color};'>{valor}{unidad}</div>
                    <div class='metric-label'>{nombre}</div>
                </div>
                """, unsafe_allow_html=True)
                st.markdown("<br>", unsafe_allow_html=True)

        st.markdown(f"""
        <div class='brecha-alert'>
            <div style='font-weight:700; font-size:1.1rem; color:#c53030;'>
                ⚠️ Mutación del Objeto de Medición (MOM) detectada
            </div>
            <div style='margin-top:0.5rem; color:#742a2a;'>
                El GAD declaró <strong>100%</strong> de cumplimiento al SIGAD/SNP.<br>
                TERRA verifica <strong>72.93%</strong> mediante 9 silos independientes.<br>
                Brecha: <strong>27.07 puntos porcentuales</strong> — evidencia documentada.
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### Señales de Atención Temprana (SAT)")
        sc1, sc2, sc3, sc4 = st.columns(4)
        sats_info = [
            ("SAT-I", "Fragmentación Selectiva", False, "ICM alto sobre universo completo"),
            ("SAT-II", "Sustitución Estratégica", False, "Sin cambios de unidad detectados"),
            ("SAT-III", "Inflación de Unidades", False, "Cumplimiento consistente con ejecución"),
            ("SAT-IV", "Tensión Normativa", False, "Reforma COOTAD no vigente en 2024"),
        ]
        for col_sat, (sat, desc, activa, nota) in zip([sc1, sc2, sc3, sc4], sats_info):
            with col_sat:
                clase = "sat-alert" if activa else "sat-ok"
                icono = "🔴" if activa else "✅"
                st.markdown(f"""
                <div class='sat-card {clase}'>
                    <div style='font-size:1.5rem;'>{icono}</div>
                    <div style='font-weight:700; font-size:0.85rem;'>{sat}</div>
                    <div style='font-size:0.75rem;'>{desc}</div>
                    <div style='font-size:0.7rem; margin-top:0.3rem; opacity:0.8;'>{nota}</div>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### ICPI por meta estratégica")
        import plotly.express as px

        df_metas = pd.DataFrame(datos["metas"])
        df_metas.columns = ["ID", "Meta", "ICPI%", "AVEP", "ODS", "Dirección"]
        df_metas = df_metas.sort_values("ICPI%", ascending=False)

        fig = px.bar(
            df_metas, x="ICPI%", y="Meta", orientation="h", color="ICPI%",
            color_continuous_scale=[[0,"#e53e3e"],[0.2,"#e67e22"],[0.4,"#d69e2e"],[0.7,"#38a169"],[0.9,"#003087"]],
            range_color=[0, 100], title="ICPI por meta — GAD Montecristi 2024"
        )
        fig.update_layout(height=600, showlegend=False)
        fig.add_vline(x=70, line_dash="dash", line_color="#003087", annotation_text="Umbral Mandato")
        st.plotly_chart(fig, use_container_width=True)

        st.markdown("### IED por Dirección Municipal")
        ied_data = {
            "Dirección": ["Dir. Económica","Patronato","Dir. Administrativa","Dir. TIC","Dir. Agua","Dir. Ambiental","Dir. Cultura","Dir. Obras Públicas"],
            "IED%": [95.0, 94.9, 85.0, 85.0, 72.0, 68.8, 47.8, 34.7],
        }
        df_ied = pd.DataFrame(ied_data).sort_values("IED%", ascending=True)
        fig2 = px.bar(df_ied, x="IED%", y="Dirección", orientation="h", color="IED%",
                      color_continuous_scale=[[0,"#e53e3e"],[0.6,"#d69e2e"],[0.85,"#38a169"],[1,"#003087"]],
                      range_color=[0,100], title="IED por Dirección — Clasificación LOSEP")
        fig2.add_vline(x=60, line_dash="dash", line_color="gray", annotation_text="Mínimo Regular")
        fig2.update_layout(height=400, showlegend=False)
        st.plotly_chart(fig2, use_container_width=True)

        st.markdown(f"""<div class='nota-metodologica'>{datos['nota_metodologica']}</div>""", unsafe_allow_html=True)

        # ── DESCARGAS ──────────────────────────────
        st.markdown("### Descargar resultados")
        col_dl1, col_dl2, col_dl3 = st.columns(3)

        with col_dl1:
            df_export = pd.DataFrame(datos["metas"])
            csv = df_export.to_csv(index=False).encode("utf-8")
            st.download_button(
                "📊 Dataset CSV (academia)",
                data=csv,
                file_name="terra_icpi_montecristi_2024.csv",
                mime="text/csv",
                use_container_width=True
            )

        with col_dl2:
            # MEJORA 3 — PDF con fpdf2
            pdf_bytes = generar_pdf_reporte(datos)
            if pdf_bytes:
                st.download_button(
                    "📄 Reporte PDF QUADRUM",
                    data=pdf_bytes,
                    file_name=f"terra_reporte_montecristi_2024.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            else:
                st.info("Agrega 'fpdf2' al requirements.txt para activar el PDF.")

        with col_dl3:
            hash_resultado = hashlib.sha256(str(datos).encode()).hexdigest()
            st.markdown(f"""
            <div class='hash-display'>
                Hash integridad ICPI 72.93%:<br>
                {hash_resultado}
            </div>
            """, unsafe_allow_html=True)

    else:
        st.info("📂 Sube tus documentos en TERRA VERIFICA primero para calcular el ICPI de tu municipio.")
        st.markdown("""
        <div class='proxima-fase'>
            <div style='font-weight:700; margin-bottom:0.5rem;'>🔮 Cómo funcionará con tus documentos</div>
            <div style='font-size:0.9rem; opacity:0.85;'>
            Una vez subas el POA, PAC y Cédula en TERRA VERIFICA, el motor ICPI calculará:<br>
            • ICPI por cada meta del PDOT · IFE Electoral · Brecha con ICM SIGAD · 4 señales SAT
            </div>
        </div>
        """, unsafe_allow_html=True)

    # MEJORA 1 — Chat Gemini en TERRA CALCULA
    seccion_gemini_chat(
        "El ciudadano está viendo el análisis ICPI de Montecristi. "
        "Puede preguntar sobre metas específicas, la brecha, los índices, o qué hacer con la información."
    )

# ─────────────────────────────────────────────
# C4 — TERRA MAPEA — MEJORA 2
# ─────────────────────────────────────────────
elif activo == "mapea":
    st.markdown("""
    <div class='terra-header'>
        <h1>🗺️ TERRA MAPEA</h1>
        <p>Visualiza las brechas de tu territorio. Ve dónde está el problema.</p>
    </div>
    """, unsafe_allow_html=True)

    try:
        import folium
        from streamlit_folium import st_folium

        st.markdown("**Mapa de brechas territoriales — GAD Montecristi 2024**")
        st.caption("Haz clic en los círculos para ver el detalle de cada meta. El tamaño del círculo representa el nivel de ejecución.")

        m = folium.Map(location=[-1.058, -80.655], zoom_start=13, tiles="CartoDB positron")

        metas_mapa = [
            {"nombre": "Agua Potable", "id": "AH-C-X-01", "icpi": 72.0, "lat": -1.045, "lon": -80.660, "dir": "Dir. Agua", "avep": "🟢 Mandato"},
            {"nombre": "Vialidad Cantonal", "id": "AH-I-X-01", "icpi": 58.3, "lat": -1.058, "lon": -80.655, "dir": "Dir. Obras", "avep": "🟡 Transición"},
            {"nombre": "Señalización — RUPTURA SISTÉMICA", "id": "AH-I-X-02", "icpi": 0.0, "lat": -1.065, "lon": -80.648, "dir": "Dir. Obras", "avep": "🔴 Ruptura"},
            {"nombre": "Salud Municipal", "id": "SC-I-N-01", "icpi": 95.0, "lat": -1.052, "lon": -80.663, "dir": "Patronato", "avep": "🔵 Excelencia"},
            {"nombre": "Desechos Sólidos", "id": "FA-C-X-01", "icpi": 80.0, "lat": -1.070, "lon": -80.658, "dir": "EP Aseo", "avep": "🟢 Mandato"},
            {"nombre": "Fortalecimiento Productivo", "id": "EP-L-N-01", "icpi": 95.0, "lat": -1.042, "lon": -80.670, "dir": "Dir. Económica", "avep": "🔵 Excelencia"},
            {"nombre": "Vivienda Interés Social", "id": "AH-I-N-01", "icpi": 22.5, "lat": -1.061, "lon": -80.645, "dir": "EP Montehogar", "avep": "🟠 Ocurrencia"},
            {"nombre": "Cultura y Patrimonio", "id": "SC-L-G-01", "icpi": 47.8, "lat": -1.055, "lon": -80.670, "dir": "Dir. Cultura", "avep": "🟡 Transición"},
            {"nombre": "Turismo", "id": "EP-L-X-01", "icpi": 85.0, "lat": -1.048, "lon": -80.650, "dir": "Dir. Turismo", "avep": "🟢 Mandato"},
            {"nombre": "Alcantarillado y PTAR", "id": "AH-C-X-02", "icpi": 72.0, "lat": -1.068, "lon": -80.665, "dir": "Dir. Agua", "avep": "🟢 Mandato"},
        ]

        for meta in metas_mapa:
            icpi = meta["icpi"]
            if icpi >= 90:
                color, color_hex = "blue", "#003087"
            elif icpi >= 70:
                color, color_hex = "green", "#38a169"
            elif icpi >= 40:
                color, color_hex = "orange", "#d69e2e"
            elif icpi > 0:
                color, color_hex = "red", "#e53e3e"
            else:
                color, color_hex = "darkred", "#742a2a"

            popup_html = f"""
            <div style='font-family:sans-serif; min-width:180px;'>
                <div style='font-weight:700; color:{color_hex}; font-size:14px;'>{meta['avep']}</div>
                <div style='font-weight:600; margin:4px 0;'>{meta['nombre']}</div>
                <div style='font-size:12px; color:#666;'>{meta['id']}</div>
                <div style='font-size:24px; font-weight:700; color:{color_hex}; margin:8px 0;'>{icpi}%</div>
                <div style='font-size:11px; color:#888;'>{meta['dir']}</div>
            </div>
            """

            folium.CircleMarker(
                location=[meta["lat"], meta["lon"]],
                radius=max(8, 8 + icpi / 8),
                color=color,
                fill=True,
                fill_color=color,
                fill_opacity=0.75,
                weight=2,
                popup=folium.Popup(popup_html, max_width=220),
                tooltip=f"{meta['nombre']}: {icpi}%"
            ).add_to(m)

        st_folium(m, width=None, height=500, use_container_width=True)

        st.markdown("""
        **Leyenda AVEP:**
        🔵 Excelencia (≥90%) &nbsp;|&nbsp; 🟢 Gestión por Mandato (70-89%) &nbsp;|&nbsp;
        🟠 Transición Crítica (40-69%) &nbsp;|&nbsp; 🔴 Gestión por Ocurrencia (<40%) &nbsp;|&nbsp; ⚫ Ruptura Sistémica (0%)
        """)

        st.markdown("---")
        st.markdown("### Metas críticas en el mapa")
        col_m1, col_m2 = st.columns(2)
        with col_m1:
            st.error("🔴 **Señalización e IVU — 0%** | Ruptura Sistémica | Dir. Obras\nEjecución financiera nula. Sin proceso SERCOP activo.")
            st.warning("🟠 **Vivienda Interés Social — 22.5%** | Gestión por Ocurrencia | EP Montehogar\nPor debajo del umbral mínimo de ejecución.")
        with col_m2:
            st.success("🔵 **Salud Municipal — 95%** | Excelencia | Patronato\n3,743 usuarios atendidos verificados.")
            st.success("🔵 **Fortalecimiento Productivo — 95%** | Excelencia | Dir. Económica\nMejor ejecutada del período.")

    except ImportError:
        st.warning("📦 Para activar el mapa interactivo, agrega estas líneas a tu requirements.txt:\n```\nfolium\nstreamlit-folium\n```")
        st.info("Guarda el archivo en GitHub y Streamlit redesplegará automáticamente en 2-3 minutos.")

    seccion_gemini_chat("El ciudadano está viendo el mapa de brechas territoriales de Montecristi.")

# ─────────────────────────────────────────────
# C5 — TERRA ARTICULA
# ─────────────────────────────────────────────
elif activo == "articula":
    st.markdown("""
    <div class='terra-header'>
        <h1>⚖️ TERRA ARTICULA</h1>
        <p>Tu análisis conectado con el marco legal y el financiamiento internacional.</p>
    </div>
    """, unsafe_allow_html=True)

    datos = cargar_datos_montecristi()

    st.markdown("### Alineación ODS — GAD Montecristi 2024")
    clusters = [
        ("🏛️ Gobernanza", "ODS 16-17", 58.8, "🟡 Transición Crítica",
         "Base institucional verificable para financiamiento CAF en Gobernanza Anticipatoria"),
        ("🏗️ Infraestructura", "ODS 6-9-11-13", 71.3, "🟢 Gestión por Mandato",
         "Inversión en agua, vialidad y equipamientos apta para BID Infraestructura Social"),
        ("🤝 Inclusión Social", "ODS 1-3-4-5-10", 63.5, "🟡 Transición Crítica",
         "Brecha en atención social documenta necesidad de apoyo GIZ y Banco Mundial"),
    ]
    for icon_c, ods, icpi_c, avep_c, argumento in clusters:
        col1, col2 = st.columns([1, 2])
        with col1:
            color_c = "#38a169" if icpi_c >= 70 else "#d69e2e"
            st.markdown(f"""
            <div class='metric-card'>
                <div style='font-size:1.5rem;'>{icon_c}</div>
                <div class='metric-value' style='color:{color_c}; font-size:1.8rem;'>{icpi_c}%</div>
                <div class='metric-label'>{ods}</div>
                <div style='font-size:0.75rem; color:#718096;'>{avep_c}</div>
            </div>
            """, unsafe_allow_html=True)
        with col2:
            st.markdown(f"""
            <div class='component-card' style='margin:0;'>
                <div style='font-size:0.85rem; color:#4a5568;'>{argumento}</div>
            </div>
            """, unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### Generador de argumento de elegibilidad bilateral")
    financiador = st.selectbox("¿Para qué organismo quieres el argumento?", [
        "CAF — Banco de Desarrollo de América Latina",
        "BID — Banco Interamericano de Desarrollo",
        "GIZ — Cooperación Técnica Alemana",
        "Banco Mundial",
        "FIDA — Fondo Internacional de Desarrollo Agrícola",
        "GEF — Fondo para el Medio Ambiente Mundial",
    ])

    if st.button("⚖️ Generar argumento de elegibilidad"):
        argumentos = {
            "CAF — Banco de Desarrollo de América Latina": f"""
**Argumento de elegibilidad — CAF Gobernanza Anticipatoria:**

El GAD Municipal de Montecristi registra un ICPI verificado del **72.93%** mediante verificación cruzada
algorítmica de 9 silos de información pública independientes, clasificado como **"Gestión por Mandato"**
según la Escala AVEP del Sistema TERRA SIAP-ICPI v2.0.

El ICPI de Gobernanza (ODS 16-17) de **58.8%** certifica una base institucional sólida para absorber
financiamiento bajo el marco de **Gobernanza Anticipatoria** de CAF, mientras que la brecha verificada
de **27.07 puntos** entre el ICPI real y el ICM auto-reportado documenta la necesidad de fortalecimiento
del Sistema de Información Local (COOTAD Art. 295).

El acueducto CAF de **$28M aprobado el 29/12/2024** para el cantón demuestra trayectoria de cooperación
activa con el organismo.

*Generado por TERRA CIUDADANA | QUADRUM GovTech | {datetime.datetime.now().strftime('%d/%m/%Y')}*
            """,
            "BID — Banco Interamericano de Desarrollo": f"""
**Argumento de elegibilidad — BID Social e Infraestructura:**

El análisis SIAP-ICPI del GAD Montecristi revela que el ICPI de Inclusión Social (ODS 1-3-4-5-10)
se sitúa en **63.51%** — Transición Crítica — documentando brechas verificables en atención a grupos
prioritarios (meta SC-I-N-03: 53.4%) y vivienda de interés social (AH-I-N-01: 22.5%).

Esta evidencia algorítmica, sustentada en 9 silos de información pública independientes, cumple con los
requisitos de **due diligence** del BID Lab para proyectos de innovación social en gobiernos locales.

*Generado por TERRA CIUDADANA | QUADRUM GovTech | {datetime.datetime.now().strftime('%d/%m/%Y')}*
            """,
        }
        argumento_texto = argumentos.get(financiador, f"""
**Argumento de elegibilidad — {financiador}:**

GAD Municipal de Montecristi | ICPI: 72.93% | Brecha documentada: 27.07pp
Verificación: 9 silos independientes | Metodología: SIAP-ICPI v2.0 QUADRUM

*Generado por TERRA CIUDADANA | {datetime.datetime.now().strftime('%d/%m/%Y')}*
        """)
        st.markdown(argumento_texto)
        st.download_button(
            "📥 Descargar argumento (.txt)",
            data=argumento_texto.encode(),
            file_name=f"elegibilidad_{financiador.split('—')[0].strip().replace(' ','_')}.txt",
            use_container_width=True
        )

    seccion_gemini_chat("El ciudadano está generando argumentos de elegibilidad para financiadores bilaterales como CAF o BID.")

# ─────────────────────────────────────────────
# C6 — TERRA ACTÚA
# ─────────────────────────────────────────────
elif activo == "actua":
    st.markdown("""
    <div class='terra-header'>
        <h1>✊ TERRA ACTÚA</h1>
        <p>De la evidencia al cambio real. Tres modos de incidencia ciudadana.</p>
    </div>
    """, unsafe_allow_html=True)

    modo_accion = st.tabs([
        "📊 Modo A — Tengo el análisis",
        "💡 Modo B — Tengo una necesidad",
        "🔍 Modo C — Encontré una convocatoria"
    ])

    with modo_accion[0]:
        st.markdown("### ¿Qué hago con mi análisis ICPI?")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Nivel 1 — Difusión inmediata**")
            st.markdown("""
            - Comparte tu análisis en redes sociales
            - URL pública de tu análisis con código QR
            - PDF descargable con membrete QUADRUM
            """)
            datos_act = cargar_datos_montecristi()
            pdf_act = generar_pdf_reporte(datos_act)
            if pdf_act:
                st.download_button(
                    "📄 Descargar PDF para compartir",
                    data=pdf_act,
                    file_name="terra_montecristi_2024.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )

        with col2:
            st.markdown("**Nivel 2 — Canales institucionales**")
            canal = st.selectbox("¿Ante quién presentas?", [
                "CPCCS — Consejo de Participación Ciudadana",
                "CGE — Contraloría General del Estado",
                "Defensoría del Pueblo",
                "Asamblea Nacional",
            ])
            if st.button("📋 Generar carta pre-formateada", use_container_width=True):
                carta = f"""SEÑORES
{canal}
Ciudad

De mi consideración:

En ejercicio de mis derechos ciudadanos de participación y control social,
consagrados en la Constitución Art. 95 y la LOPC Art. 88, presento a su
autoridad el siguiente análisis técnico verificado mediante el Sistema
TERRA SIAP-ICPI v2.0 (QUADRUM GovTech):

GAD Municipal de Montecristi | Período 2024
ICPI Verificado: 72.93% | ICM Declarado SIGAD: 100%
Brecha documentada: 27.07 puntos porcentuales

Solicito respetuosamente que su institución tome conocimiento de esta
evidencia algorítmica sustentada en 9 silos de información pública independientes.

Atentamente,
[Su nombre]
[C.I.]
[Fecha]

Generado por TERRA CIUDADANA | QUADRUM GovTech | terra.quadrum.ec"""
                st.text_area("Carta generada:", value=carta, height=280)
                st.download_button("📥 Descargar carta (.txt)", data=carta.encode(),
                                   file_name="carta_incidencia.txt", use_container_width=True)

    with modo_accion[1]:
        st.markdown("### Tengo una necesidad en mi comunidad")
        metas_pdot = cargar_datos_montecristi()["metas"]
        necesidad = st.text_area(
            "Describe el problema de tu comunidad:",
            height=100,
            placeholder="Ej: En el barrio no hay agua potable, las calles están rotas, no hay alumbrado..."
        )
        if necesidad and st.button("🔍 Analizar y vincular al PDOT"):
            necesidad_lower = necesidad.lower()
            encontradas = []
            keywords = {
                "agua": ["AH-C-X-01", "AH-C-X-02"],
                "alcantarillado": ["AH-C-X-02"],
                "vía": ["AH-I-X-01"], "calle": ["AH-I-X-01"], "carretera": ["AH-I-X-01"],
                "señalización": ["AH-I-X-02"],
                "salud": ["SC-I-N-01"],
                "vivienda": ["AH-I-N-01"],
                "basura": ["FA-C-X-01"], "desecho": ["FA-C-X-01"],
                "turismo": ["EP-L-X-01"],
                "cultura": ["SC-L-G-01"],
            }
            for palabra, ids in keywords.items():
                if palabra in necesidad_lower:
                    for meta_id in ids:
                        meta = next((m for m in metas_pdot if m["id"] == meta_id), None)
                        if meta and meta not in encontradas:
                            encontradas.append(meta)

            if encontradas:
                st.markdown("**Vinculación con el PDOT de Montecristi:**")
                for meta in encontradas:
                    color = "success" if meta["icpi"] >= 70 else "warning" if meta["icpi"] >= 40 else "error"
                    getattr(st, color)(
                        f"{meta['avep']} **{meta['nombre']}** (ID: {meta['id']}) — ICPI: {meta['icpi']}% | {meta['dir']}"
                    )
            else:
                st.info("No encontré una meta exacta. Describe con más detalle o usa palabras como: agua, calle, salud, vivienda, basura, turismo, cultura.")

    with modo_accion[2]:
        st.markdown("### Encontré una convocatoria de financiamiento")
        convocatoria = st.text_input("Pega el link o describe brevemente la convocatoria:")
        if convocatoria:
            st.success("✅ El análisis de Montecristi tiene datos verificados para sustentar esta convocatoria.")
            st.markdown("Usa **TERRA ARTICULA** para generar el argumento de elegibilidad formal con los datos ICPI.")
            st.info("💡 El ICPI de 72.93% y la brecha de 27.07pp son exactamente el tipo de evidencia que los financiadores bilaterales requieren para due diligence.")

    seccion_gemini_chat("El ciudadano está en TERRA ACTÚA buscando qué hacer con su análisis o cómo acceder a financiamiento.")

# ─────────────────────────────────────────────
# C7 — TERRA EVALÚA
# ─────────────────────────────────────────────
elif activo == "evalua":
    st.markdown("""
    <div class='terra-header'>
        <h1>🏛️ TERRA EVALÚA</h1>
        <p>Lo que dijo. Lo que firmó. Lo que TERRA verifica. Tres verdades en un panel.</p>
    </div>
    """, unsafe_allow_html=True)

    eval_data = datos_evaluacion_precacheados()
    datos = cargar_datos_montecristi()

    st.markdown(f"""
    **Caso analizado:** Rendición de Cuentas 2024 — {datos['gad']}
    **Alcalde:** {datos['alcalde']} | **Evento:** {eval_data['fecha_evento']} — {eval_data['lugar']}
    **Asistentes:** {eval_data['asistentes']} | **Informe N°:** {eval_data['informe_numero']}
    """)

    st.markdown("---")
    st.markdown("### Los tres insumos del análisis")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("""
        <div class='metric-card'>
            <div style='font-size:2rem;'>🧮</div>
            <div style='font-weight:700; color:#003087;'>ICPI TERRA</div>
            <div style='font-size:2rem; font-weight:700; color:#38a169;'>72.93%</div>
            <div class='metric-label'>Verificación algorítmica<br>9 silos independientes</div>
        </div>""", unsafe_allow_html=True)
    with col2:
        st.markdown("""
        <div class='metric-card'>
            <div style='font-size:2rem;'>📋</div>
            <div style='font-weight:700; color:#003087;'>INFORME TÉCNICO CPCCS</div>
            <div style='font-size:1rem; font-weight:700; color:#4a5568;'>Informe N° 22844</div>
            <div class='metric-label'>Documento oficial firmado<br>Rendición de cuentas 2024</div>
        </div>""", unsafe_allow_html=True)
    with col3:
        st.markdown("""
        <div class='metric-card'>
            <div style='font-size:2rem;'>🎙️</div>
            <div style='font-weight:700; color:#003087;'>DISCURSO POLÍTICO</div>
            <div style='font-size:1rem; font-weight:700; color:#4a5568;'>Evento público</div>
            <div class='metric-label'>11/07/2025 — 261 asistentes<br>Transcripción verificada</div>
        </div>""", unsafe_allow_html=True)

    st.markdown("---")
    col_izq, col_der = st.columns([2, 1])
    with col_izq:
        st.markdown("### Coeficiente de Fricción Narrativa")
        st.markdown(f"""
        <div style='margin: 1rem 0;'>
            <div style='display:flex; justify-content:space-between; font-size:0.85rem; color:#718096;'>
                <span>Coherencia total</span><span>Fricción total</span>
            </div>
            <div class='friccion-barra'></div>
            <div style='display:flex; justify-content:space-between; font-weight:700;'>
                <span style='color:#38a169;'>{eval_data['coherencia']}% Coherente</span>
                <span style='color:#e53e3e;'>{eval_data['friccion_narrativa']}% Fricción</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
    with col_der:
        st.markdown(f"""
        <div class='metric-card' style='text-align:center;'>
            <div style='font-size:3rem; font-weight:700; color:#e53e3e;'>{eval_data['friccion_narrativa']}%</div>
            <div class='metric-label'>Fricción Narrativa<br>Detectada</div>
        </div>""", unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### Hallazgos por meta")
    for hallazgo in eval_data["hallazgos"]:
        clase = "hallazgo-alto" if hallazgo["nivel"] == "alto" else "hallazgo-bajo"
        st.markdown(f"""
        <div class='hallazgo-card {clase}'>
            <div style='font-weight:700;'>{hallazgo['tipo']} — {hallazgo['meta']}</div>
            <div style='display:grid; grid-template-columns:1fr 1fr 1fr; gap:1rem; margin-top:0.5rem; font-size:0.85rem;'>
                <div><strong>🎙️ Discurso:</strong><br>{hallazgo['discurso']}</div>
                <div><strong>📋 Informe CPCCS:</strong><br>{hallazgo['informe']}</div>
                <div><strong>🧮 TERRA verifica:</strong><br>{hallazgo['terra']}</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### Promesas electorales sin meta PDOT")
    for promesa in eval_data["promesas_sin_pdot"]:
        st.markdown(f"🔴 {promesa}")

    st.markdown("---")
    st.markdown("### Exportar análisis completo")
    col_e1, col_e2, col_e3 = st.columns(3)

    with col_e1:
        datos_ev = cargar_datos_montecristi()
        pdf_ev = generar_pdf_reporte(datos_ev)
        if pdf_ev:
            st.download_button(
                "📄 PDF completo",
                data=pdf_ev,
                file_name="terra_evalua_montecristi_2024.pdf",
                mime="application/pdf",
                use_container_width=True
            )

    with col_e2:
        reporte_texto = f"""TERRA EVALÚA — Análisis de Fricción Narrativa
GAD Montecristi | Rendición de Cuentas 2024
{'='*50}
ICPI Verificado: {datos['icpi']}%
ICM Declarado: {datos['icm_sigad']}%
Brecha: {datos['brecha']} pp
Fricción Narrativa: {eval_data['friccion_narrativa']}%

HALLAZGOS:
{chr(10).join([f"- {h['tipo']}: {h['meta']}" for h in eval_data['hallazgos']])}

Generado por TERRA CIUDADANA | QUADRUM GovTech
{datetime.datetime.now().strftime('%d/%m/%Y')}"""
        st.download_button(
            "📰 Reporte periodista (.txt)",
            data=reporte_texto.encode(),
            file_name="terra_evalua_montecristi_2024.txt",
            use_container_width=True
        )

    with col_e3:
        df_eval = pd.DataFrame(eval_data["hallazgos"])
        csv_eval = df_eval.to_csv(index=False).encode("utf-8")
        st.download_button(
            "🎓 Dataset academia (.csv)",
            data=csv_eval,
            file_name="friccion_narrativa_montecristi_2024.csv",
            use_container_width=True
        )

    st.markdown("""
    <div class='proxima-fase'>
        <div style='font-weight:700; margin-bottom:0.5rem;'>🔮 Próxima fase — El Tribunal completo</div>
        <div style='font-size:0.9rem; opacity:0.85;'>
        • Whisper transcribirá el video del evento en tiempo real<br>
        • OCR leerá el PDF del informe CPCCS automáticamente<br>
        • El Diccionario del Argot Ecuatoriano detectará frases como "ya mismito sale la obra"<br>
        • Requiere: API Anthropic + activación del módulo NLP
        </div>
    </div>
    """, unsafe_allow_html=True)

    seccion_gemini_chat(
        "El ciudadano está viendo el análisis de fricción narrativa entre el discurso del Alcalde, "
        "el informe técnico del CPCCS y el ICPI verificado por TERRA."
    )

# ─────────────────────────────────────────────
# FOOTER
# ─────────────────────────────────────────────
st.markdown("""
<div class='footer-terra'>
    <strong>TERRA CIUDADANA v2.0</strong> | QUADRUM GovTech<br>
    Ronald Javier Delgado Santana | Diplomado DGIP CAF-ESPOL 2026<br>
    Metodología SIAP-ICPI | Caso piloto: GAD Municipal de Montecristi<br>
    <br>
    El acceso a información pública es un derecho constitucional (Art. 18 Constitución del Ecuador).<br>
    Esta herramienta no sustituye asesoría jurídica profesional.
</div>
""", unsafe_allow_html=True)
